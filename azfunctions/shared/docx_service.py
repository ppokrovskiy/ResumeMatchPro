import io
from docx import Document


class DocxService:
    def get_text_from_docx(document_content):
        doc = Document(io.BytesIO(document_content))
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        # extract text also from tables and other elements
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    fullText.append(cell.text)
        # extract text from text boxes
        for shape in doc.inline_shapes:
            if shape.text_frame:
                for paragraph in shape.text_frame.paragraphs:
                    fullText.append(paragraph.text)
        # extract text from shapes
        if hasattr(doc, 'shapes'):
            for shape in doc.shapes:
                if shape.has_text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        fullText.append(paragraph.text)
        # extract text from headers
        # for section in doc.sections:
        #     fullText.append(section.header.text)
        # # extract text from footers
        # for section in doc.sections:
        #     for footer in section.footer:
        #         fullText.append(footer.text)
        return '\n'.join(fullText)

    # Replace 'your_document.docx' with the path to your Word document
    # document_path = 'your_document.docx'
    # document_text = get_text_from_docx(document_path)
    # print(document_text)
